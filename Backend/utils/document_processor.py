import PyPDF2
import docx
import io
import logging
import re
from typing import Dict, List, Optional, Tuple
from werkzeug.datastructures import FileStorage

logger = logging.getLogger(__name__)

# Optional imports with fallbacks
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False
    logger.warning("pdfplumber not available, using PyPDF2 only for PDF processing")

try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False
    logger.warning("python-magic not available, using file extensions for type detection")

try:
    import pytesseract
    from PIL import Image
    import fitz  # PyMuPDF
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    logger.warning("OCR dependencies not available (pytesseract, PIL, PyMuPDF)")

class DocumentProcessor:
    """Enhanced document processing utilities for extracting text from various file formats."""
    
    SUPPORTED_FORMATS = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/msword': 'doc',
        'text/plain': 'txt'
    }
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common PDF artifacts
        text = re.sub(r'[^\w\s\.,!?;:()\[\]{}"\'-]', ' ', text)
        
        # Fix common OCR errors
        text = re.sub(r'\b[A-Z]{2,}\b', lambda m: m.group().title(), text)  # Fix all caps words
        
        # Remove excessive line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    @staticmethod
    def extract_metadata_summary(metadata: Dict) -> str:
        """
        Create a summary string from document metadata.
        
        Args:
            metadata: Document metadata dictionary
            
        Returns:
            Formatted metadata summary
        """
        summary_parts = []
        
        if metadata.get('title'):
            summary_parts.append(f"Title: {metadata['title']}")
        if metadata.get('author'):
            summary_parts.append(f"Author: {metadata['author']}")
        if metadata.get('subject'):
            summary_parts.append(f"Subject: {metadata['subject']}")
        if metadata.get('page_count'):
            summary_parts.append(f"Pages: {metadata['page_count']}")
        if metadata.get('word_count'):
            summary_parts.append(f"Words: {metadata['word_count']}")
        
        return " | ".join(summary_parts) if summary_parts else "No metadata available"
    
    @staticmethod
    def validate_file(file: FileStorage, max_size: int = 10485760) -> Tuple[bool, str]:
        """
        Validate uploaded file.
        
        Args:
            file: Uploaded file object
            max_size: Maximum file size in bytes (default 10MB)
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not file:
            return False, "No file provided"
        
        if file.filename == '':
            return False, "No file selected"
        
        # Check file size
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > max_size:
            return False, f"File size ({file_size} bytes) exceeds maximum allowed ({max_size} bytes)"
        
        if file_size == 0:
            return False, "File is empty"
        
        return True, ""
    
    @staticmethod
    def detect_file_type(file: FileStorage) -> Optional[str]:
        """
        Detect file type using python-magic or fallback to filename.
        
        Args:
            file: Uploaded file object
            
        Returns:
            Detected MIME type or None
        """
        try:
            if HAS_MAGIC:
                file.seek(0)
                file_content = file.read(1024)  # Read first 1KB
                file.seek(0)  # Reset
                
                mime_type = magic.from_buffer(file_content, mime=True)
                return mime_type
            else:
                # Fallback to filename extension
                if file.filename:
                    filename = file.filename.lower()
                    if filename.endswith('.pdf'):
                        return 'application/pdf'
                    elif filename.endswith('.docx'):
                        return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    elif filename.endswith('.doc'):
                        return 'application/msword'
                    elif filename.endswith('.txt'):
                        return 'text/plain'
                return None
        except Exception as e:
            logger.error(f"Error detecting file type: {str(e)}")
            return None
    
    @staticmethod
    def extract_text_from_pdf(file: FileStorage) -> Dict[str, any]:
        """
        Extract text from PDF file using multiple methods including OCR fallback.
        
        Args:
            file: PDF file object
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            file.seek(0)
            result = {
                'text': '',
                'pages': [],
                'metadata': {},
                'method': 'pdfplumber',
                'success': True,
                'error': None,
                'ocr_used': False
            }
            
            # Try pdfplumber first (better for complex layouts) if available
            if HAS_PDFPLUMBER:
                try:
                    with pdfplumber.open(file) as pdf:
                        result['metadata'] = {
                            'page_count': len(pdf.pages),
                            'title': pdf.metadata.get('Title', ''),
                            'author': pdf.metadata.get('Author', ''),
                            'subject': pdf.metadata.get('Subject', ''),
                            'creator': pdf.metadata.get('Creator', ''),
                            'creation_date': str(pdf.metadata.get('CreationDate', '')),
                            'modification_date': str(pdf.metadata.get('ModDate', ''))
                        }
                        
                        full_text = ""
                        pages_text = []
                        pages_with_little_text = []
                        
                        for page_num, page in enumerate(pdf.pages, 1):
                            try:
                                page_text = page.extract_text() or ""
                                cleaned_text = DocumentProcessor.clean_text(page_text)
                                
                                pages_text.append({
                                    'page_number': page_num,
                                    'text': cleaned_text,
                                    'char_count': len(cleaned_text),
                                    'word_count': len(cleaned_text.split()) if cleaned_text else 0
                                })
                                
                                # Check if page has very little text (might be image-based)
                                if len(cleaned_text.strip()) < 50:
                                    pages_with_little_text.append(page_num)
                                
                                full_text += f"\n--- Page {page_num} ---\n{cleaned_text}\n"
                                
                            except Exception as e:
                                logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                                pages_text.append({
                                    'page_number': page_num,
                                    'text': '',
                                    'char_count': 0,
                                    'word_count': 0,
                                    'error': str(e)
                                })
                                pages_with_little_text.append(page_num)
                        
                        # Try OCR on pages with little text if OCR is available
                        if HAS_OCR and pages_with_little_text and len(pages_with_little_text) < len(pdf.pages) * 0.8:
                            logger.info(f"Attempting OCR on {len(pages_with_little_text)} pages with little text")
                            ocr_text = DocumentProcessor._try_ocr_on_pdf_pages(file, pages_with_little_text)
                            if ocr_text:
                                result['ocr_used'] = True
                                full_text += "\n--- OCR Results ---\n" + ocr_text
                        
                        result['text'] = DocumentProcessor.clean_text(full_text)
                        result['pages'] = pages_text
                        result['metadata']['total_word_count'] = len(result['text'].split()) if result['text'] else 0
                        result['metadata']['pages_with_little_text'] = len(pages_with_little_text)
                        
                        return result
                        
                except Exception as e:
                    logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
            
            # Use PyPDF2 (fallback or primary if pdfplumber not available)
            file.seek(0)
            
            pdf_reader = PyPDF2.PdfReader(file)
            result['method'] = 'PyPDF2'
            result['metadata'] = {
                'page_count': len(pdf_reader.pages),
                'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                'subject': pdf_reader.metadata.get('/Subject', '') if pdf_reader.metadata else ''
            }
            
            full_text = ""
            pages_text = []
            pages_with_little_text = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text() or ""
                    cleaned_text = DocumentProcessor.clean_text(page_text)
                    
                    pages_text.append({
                        'page_number': page_num,
                        'text': cleaned_text,
                        'char_count': len(cleaned_text),
                        'word_count': len(cleaned_text.split()) if cleaned_text else 0
                    })
                    
                    if len(cleaned_text.strip()) < 50:
                        pages_with_little_text.append(page_num)
                    
                    full_text += f"\n--- Page {page_num} ---\n{cleaned_text}\n"
                    
                except Exception as page_error:
                    logger.warning(f"Error extracting text from page {page_num}: {str(page_error)}")
                    pages_text.append({
                        'page_number': page_num,
                        'text': '',
                        'char_count': 0,
                        'word_count': 0,
                        'error': str(page_error)
                    })
                    pages_with_little_text.append(page_num)
            
            # Try OCR if many pages have little text and OCR is available
            if HAS_OCR and pages_with_little_text and len(pages_with_little_text) >= len(pdf_reader.pages) * 0.3:
                logger.info(f"Attempting OCR on PDF with {len(pages_with_little_text)} pages with little text")
                ocr_text = DocumentProcessor._try_ocr_on_pdf(file)
                if ocr_text:
                    result['ocr_used'] = True
                    full_text += "\n--- OCR Results ---\n" + ocr_text
            
            result['text'] = DocumentProcessor.clean_text(full_text)
            result['pages'] = pages_text
            result['metadata']['total_word_count'] = len(result['text'].split()) if result['text'] else 0
            result['metadata']['pages_with_little_text'] = len(pages_with_little_text)
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return {
                'text': '',
                'pages': [],
                'metadata': {},
                'method': 'failed',
                'success': False,
                'error': str(e),
                'ocr_used': False
            }
    
    @staticmethod
    def _try_ocr_on_pdf(file: FileStorage) -> str:
        """
        Attempt OCR on entire PDF using PyMuPDF and Tesseract.
        
        Args:
            file: PDF file object
            
        Returns:
            OCR extracted text or empty string if failed
        """
        if not HAS_OCR:
            return ""
        
        try:
            file.seek(0)
            pdf_bytes = file.read()
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            ocr_text = ""
            for page_num in range(len(pdf_document)):
                try:
                    page = pdf_document.load_page(page_num)
                    # Convert page to image
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    
                    # Perform OCR
                    page_text = pytesseract.image_to_string(img, config='--psm 6')
                    if page_text.strip():
                        ocr_text += f"\n--- OCR Page {page_num + 1} ---\n{page_text}\n"
                        
                except Exception as e:
                    logger.warning(f"OCR failed for page {page_num + 1}: {str(e)}")
                    continue
            
            pdf_document.close()
            return DocumentProcessor.clean_text(ocr_text)
            
        except Exception as e:
            logger.error(f"OCR processing failed: {str(e)}")
            return ""
    
    @staticmethod
    def _try_ocr_on_pdf_pages(file: FileStorage, page_numbers: List[int]) -> str:
        """
        Attempt OCR on specific PDF pages.
        
        Args:
            file: PDF file object
            page_numbers: List of page numbers to OCR (1-indexed)
            
        Returns:
            OCR extracted text or empty string if failed
        """
        if not HAS_OCR or not page_numbers:
            return ""
        
        try:
            file.seek(0)
            pdf_bytes = file.read()
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            ocr_text = ""
            for page_num in page_numbers:
                try:
                    page = pdf_document.load_page(page_num - 1)  # Convert to 0-indexed
                    # Convert page to image
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    
                    # Perform OCR
                    page_text = pytesseract.image_to_string(img, config='--psm 6')
                    if page_text.strip():
                        ocr_text += f"\n--- OCR Page {page_num} ---\n{page_text}\n"
                        
                except Exception as e:
                    logger.warning(f"OCR failed for page {page_num}: {str(e)}")
                    continue
            
            pdf_document.close()
            return DocumentProcessor.clean_text(ocr_text)
            
        except Exception as e:
            logger.error(f"OCR processing failed: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file: FileStorage) -> Dict[str, any]:
        """
        Extract text from DOCX file with enhanced metadata and structure.
        
        Args:
            file: DOCX file object
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            file.seek(0)
            doc = docx.Document(file)
            
            # Extract text from paragraphs with styles
            paragraphs = []
            full_text = ""
            headings = []
            
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    # Check if this is a heading
                    style_name = para.style.name if para.style else ""
                    if "Heading" in style_name:
                        headings.append({
                            'text': para_text,
                            'style': style_name,
                            'level': self._extract_heading_level(style_name)
                        })
                        full_text += f"\n## {para_text}\n"
                    else:
                        paragraphs.append({
                            'text': para_text,
                            'style': style_name
                        })
                        full_text += para_text + "\n"
            
            # Extract text from tables with better formatting
            tables_text = []
            for table_num, table in enumerate(doc.tables, 1):
                table_data = []
                table_text = f"\n[TABLE {table_num}]\n"
                
                for row_num, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    
                    if any(row_data):  # Only add non-empty rows
                        table_data.append(row_data)
                        table_text += " | ".join(row_data) + "\n"
                
                if table_data:
                    tables_text.append({
                        'table_number': table_num,
                        'data': table_data,
                        'text': table_text
                    })
                    full_text += table_text + "\n"
            
            # Extract hyperlinks
            hyperlinks = []
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.hyperlink:
                        hyperlinks.append({
                            'text': run.text,
                            'url': run.hyperlink.address if run.hyperlink.address else ''
                        })
            
            # Get document properties
            props = doc.core_properties
            
            # Calculate statistics
            cleaned_full_text = DocumentProcessor.clean_text(full_text)
            word_count = len(cleaned_full_text.split()) if cleaned_full_text else 0
            char_count = len(cleaned_full_text)
            
            metadata = {
                'title': props.title or '',
                'author': props.author or '',
                'subject': props.subject or '',
                'description': props.comments or '',
                'keywords': props.keywords or '',
                'category': props.category or '',
                'created': str(props.created) if props.created else '',
                'modified': str(props.modified) if props.modified else '',
                'last_modified_by': props.last_modified_by or '',
                'paragraph_count': len(paragraphs),
                'table_count': len(tables_text),
                'heading_count': len(headings),
                'hyperlink_count': len(hyperlinks),
                'word_count': word_count,
                'char_count': char_count
            }
            
            return {
                'text': cleaned_full_text,
                'paragraphs': paragraphs,
                'tables': tables_text,
                'headings': headings,
                'hyperlinks': hyperlinks,
                'metadata': metadata,
                'success': True,
                'error': None
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return {
                'text': '',
                'paragraphs': [],
                'tables': [],
                'headings': [],
                'hyperlinks': [],
                'metadata': {},
                'success': False,
                'error': str(e)
            }
    
    @staticmethod
    def _extract_heading_level(style_name: str) -> int:
        """Extract heading level from style name."""
        import re
        match = re.search(r'Heading (\d+)', style_name)
        return int(match.group(1)) if match else 1
    
    @staticmethod
    def extract_text_from_txt(file: FileStorage) -> Dict[str, any]:
        """
        Extract text from TXT file.
        
        Args:
            file: TXT file object
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            file.seek(0)
            content = file.read()
            
            # Try to decode with different encodings
            encodings = ['utf-8', 'utf-16', 'ascii', 'latin-1']
            text = None
            encoding_used = None
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    encoding_used = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError("Unable to decode text file with any supported encoding")
            
            lines = text.split('\n')
            metadata = {
                'encoding': encoding_used,
                'line_count': len(lines),
                'char_count': len(text),
                'word_count': len(text.split()) if text else 0
            }
            
            return {
                'text': text,
                'lines': lines,
                'metadata': metadata,
                'success': True,
                'error': None
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            return {
                'text': '',
                'lines': [],
                'metadata': {},
                'success': False,
                'error': str(e)
            }
    
    @classmethod
    def process_document(cls, file: FileStorage) -> Dict[str, any]:
        """
        Process document and extract text based on file type with comprehensive analysis.
        
        Args:
            file: Uploaded file object
            
        Returns:
            Dictionary containing extracted text, metadata, and analysis
        """
        # Validate file
        is_valid, error_msg = cls.validate_file(file)
        if not is_valid:
            return {
                'success': False,
                'error': error_msg,
                'text': '',
                'metadata': {},
                'analysis': {}
            }
        
        # Detect file type
        mime_type = cls.detect_file_type(file)
        if mime_type not in cls.SUPPORTED_FORMATS:
            return {
                'success': False,
                'error': f"Unsupported file type: {mime_type}. Supported types: {', '.join(cls.SUPPORTED_FORMATS.values())}",
                'text': '',
                'metadata': {},
                'analysis': {}
            }
        
        file_type = cls.SUPPORTED_FORMATS[mime_type]
        
        # Process based on file type
        try:
            if file_type == 'pdf':
                result = cls.extract_text_from_pdf(file)
            elif file_type == 'docx':
                result = cls.extract_text_from_docx(file)
            elif file_type == 'txt':
                result = cls.extract_text_from_txt(file)
            else:
                return {
                    'success': False,
                    'error': f"Handler not implemented for file type: {file_type}",
                    'text': '',
                    'metadata': {},
                    'analysis': {}
                }
            
            if not result['success']:
                return result
            
            # Add common metadata
            result['metadata']['file_type'] = file_type
            result['metadata']['mime_type'] = mime_type
            result['metadata']['filename'] = file.filename
            
            # Perform text analysis
            text_analysis = cls._analyze_text(result['text'])
            result['analysis'] = text_analysis
            
            # Add quality assessment
            quality_assessment = cls._assess_extraction_quality(result)
            result['quality'] = quality_assessment
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            return {
                'success': False,
                'error': f"Error processing document: {str(e)}",
                'text': '',
                'metadata': {},
                'analysis': {}
            }
    
    @staticmethod
    def _analyze_text(text: str) -> Dict[str, any]:
        """
        Perform basic text analysis.
        
        Args:
            text: Extracted text
            
        Returns:
            Dictionary containing text analysis results
        """
        if not text:
            return {
                'word_count': 0,
                'char_count': 0,
                'sentence_count': 0,
                'paragraph_count': 0,
                'avg_words_per_sentence': 0,
                'readability_indicators': {},
                'language_characteristics': {}
            }
        
        # Basic counts
        word_count = len(text.split())
        char_count = len(text)
        
        # Sentence count (rough estimate)
        sentence_endings = text.count('.') + text.count('!') + text.count('?')
        sentence_count = max(1, sentence_endings)
        
        # Paragraph count
        paragraph_count = len([p for p in text.split('\n\n') if p.strip()])
        
        # Average words per sentence
        avg_words_per_sentence = word_count / sentence_count if sentence_count > 0 else 0
        
        # Readability indicators
        readability_indicators = {
            'avg_word_length': sum(len(word) for word in text.split()) / word_count if word_count > 0 else 0,
            'complex_word_ratio': len([w for w in text.split() if len(w) > 6]) / word_count if word_count > 0 else 0,
            'uppercase_ratio': sum(1 for c in text if c.isupper()) / char_count if char_count > 0 else 0
        }
        
        # Language characteristics
        language_characteristics = {
            'has_legal_terms': any(term in text.lower() for term in [
                'agreement', 'contract', 'clause', 'provision', 'party', 'whereas',
                'hereby', 'therefore', 'liability', 'indemnify', 'terminate'
            ]),
            'has_financial_terms': any(term in text.lower() for term in [
                'payment', 'fee', 'cost', 'price', 'amount', 'currency', 'dollar',
                'invoice', 'billing', 'compensation'
            ]),
            'has_dates': bool(re.search(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b', text)),
            'has_addresses': bool(re.search(r'\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln)\b', text, re.IGNORECASE)),
            'has_phone_numbers': bool(re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', text)),
            'has_email_addresses': bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text))
        }
        
        return {
            'word_count': word_count,
            'char_count': char_count,
            'sentence_count': sentence_count,
            'paragraph_count': paragraph_count,
            'avg_words_per_sentence': round(avg_words_per_sentence, 2),
            'readability_indicators': readability_indicators,
            'language_characteristics': language_characteristics
        }
    
    @staticmethod
    def _assess_extraction_quality(result: Dict[str, any]) -> Dict[str, any]:
        """
        Assess the quality of text extraction.
        
        Args:
            result: Extraction result dictionary
            
        Returns:
            Quality assessment dictionary
        """
        text = result.get('text', '')
        metadata = result.get('metadata', {})
        
        quality_score = 100  # Start with perfect score
        issues = []
        warnings = []
        
        # Check if text was extracted
        if not text or len(text.strip()) < 10:
            quality_score -= 50
            issues.append("Very little or no text extracted")
        
        # Check for OCR usage (lower confidence)
        if result.get('ocr_used', False):
            quality_score -= 20
            warnings.append("OCR was used - text accuracy may be lower")
        
        # Check for extraction errors
        if 'pages' in result:
            failed_pages = sum(1 for page in result['pages'] if 'error' in page)
            if failed_pages > 0:
                quality_score -= (failed_pages / len(result['pages'])) * 30
                warnings.append(f"{failed_pages} pages had extraction errors")
        
        # Check text quality indicators
        if text:
            # Too many special characters might indicate extraction issues
            special_char_ratio = len(re.findall(r'[^\w\s\.,!?;:()\[\]{}"\'-]', text)) / len(text)
            if special_char_ratio > 0.1:
                quality_score -= 15
                warnings.append("High number of special characters detected")
            
            # Very short average word length might indicate extraction issues
            words = text.split()
            if words:
                avg_word_length = sum(len(word) for word in words) / len(words)
                if avg_word_length < 3:
                    quality_score -= 10
                    warnings.append("Unusually short average word length")
        
        # Determine quality level
        if quality_score >= 90:
            quality_level = "excellent"
        elif quality_score >= 75:
            quality_level = "good"
        elif quality_score >= 60:
            quality_level = "fair"
        elif quality_score >= 40:
            quality_level = "poor"
        else:
            quality_level = "very_poor"
        
        return {
            'score': max(0, quality_score),
            'level': quality_level,
            'issues': issues,
            'warnings': warnings,
            'extraction_method': result.get('method', 'unknown'),
            'ocr_used': result.get('ocr_used', False)
        }